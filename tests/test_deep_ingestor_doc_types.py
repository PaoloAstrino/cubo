import json
from pathlib import Path
import pandas as pd
import pytest

from src.ingest.deep_ingestor import DeepIngestor


def test_docx_chunking(tmp_path: Path):
    folder = tmp_path / "docs"
    folder.mkdir()
    docx_path = folder / "doc.docx"

    # Create a docx with multiple paragraphs
    from docx import Document
    doc = Document()
    doc.add_paragraph('This is the first paragraph.')
    doc.add_paragraph('This is the second paragraph.')
    doc.add_paragraph('Third paragraph for good measure.')
    doc.save(str(docx_path))

    out = tmp_path / "out"
    res = DeepIngestor(input_folder=str(folder), output_dir=str(out)).ingest()
    df = pd.read_parquet(res['chunks_parquet'])
    # Should have sentence-window chunks with sentence index in chunk_id or sentence_index field
    assert len(df) > 0
    assert df['chunk_id'].str.contains('_s').any()


def test_excel_chunking(tmp_path: Path):
    folder = tmp_path / "docs"
    folder.mkdir()
    xlsx_path = folder / "sheets.xlsx"

    import pandas as pd
    df = pd.DataFrame({'col1': range(3), 'col2': ['a','b','c']})
    df.to_excel(xlsx_path, index=False)

    out = tmp_path / "out"
    res = DeepIngestor(input_folder=str(folder), output_dir=str(out)).ingest()
    df = pd.read_parquet(res['chunks_parquet'])
    assert 'sheet' in '\n'.join(df['file_path'].values)
    assert any([sid for sid in df['chunk_id'] if '_sheet_' in sid])


def test_pdf_chunking(tmp_path: Path):
    folder = tmp_path / "docs"
    folder.mkdir()
    pdf_path = folder / "test.pdf"

    # Create a simple PDF with two pages using reportlab
    try:
        from reportlab.pdfgen import canvas
    except Exception:
        pytest.skip('reportlab not installed - skip pdf test')

    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "Page 1 content. Hello world.")
    c.showPage()
    c.drawString(100, 750, "Page 2 content. Another page.")
    c.save()

    out = tmp_path / "out"
    res = DeepIngestor(input_folder=str(folder), output_dir=str(out)).ingest()
    df = pd.read_parquet(res['chunks_parquet'])
    assert len(df) > 0
    # PDF chunking may produce sentence windows with sentence indices
    assert any('_s' in cid for cid in df['chunk_id'])
